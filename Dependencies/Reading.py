from datetime import datetime
from docx import Document
import pandas, csv

def clean_string(line):
    dummy=line.replace(',', '').replace(';', '').replace('\n', '').split(' ') #Remove commas, semicolons and line breaks for csv format
    while '' in dummy: dummy.remove('')                                       #Remove redundant empty space
    cleaned_string=''
    for n in range(len(dummy)-1): cleaned_string+=dummy[n]+' '
    cleaned_string+=dummy[len(dummy)-1]
    return cleaned_string.lower()

def read_bank_statement(filename,san_or_rbs):    
    new_data=[]
    
    if san_or_rbs=='san':
        userfile = open(filename, 'r', encoding="latin-1") 
        line=userfile.readline()
        temp = [[],[],[]]
        while line!='':
            if line.split('\xa0')[0]=='Date:': temp[0]=line.split('\xa0')[1].split('\n')[0]
            if line.split('\xa0')[0]=='Description:': temp[1]=clean_string(line.split('\xa0')[1])
            if line.split('\xa0')[0]=='Amount:': temp[2]=line.split('\xa0')[1].split('\n')[0]
            if line.split('\xa0')[0]=='Balance:':
                new_data.append(temp)
                temp = [[],[],[]]
            line=userfile.readline()
        userfile.close()        
        
    if san_or_rbs=='rbs': #From 05/05/21 onwards, integer payments no longer have decimal places
        userfile = open(filename, 'r') 
        for dummy in range(4): line=userfile.readline()
        while line!='':
            temp=[]
            temp.append(line.split('"')[1])
            temp.append(clean_string(line.split('"')[5]))
            temp.append(line.split('"')[6].split(',')[1])
            new_data.append(temp)
            line=userfile.readline()
        userfile.close()
        
    return new_data

def read_sefton_file(filename):
    conversion=[61,28,53,13,19] #[57,25,49,13,18] before invoice batch 12 of 2018
    data=[]
    with open(filename, 'rt') as file:
        reader = csv.reader(file, quotechar='"',delimiter=',',quoting=csv.QUOTE_ALL,skipinitialspace=True)
        for row in reader: data.append(row)
    file.close()
    
    print(data[1][56]+'\n') #Print remittance note on the sefton file (usually there isn't one, but see April 2018 for example)
    true_data=[]
    for n in range(len(data)):
        if data[n][conversion[4]]=='Income': true_data.append([data[n][conversion[0]],data[n][conversion[1]],data[n][conversion[2]],data[n][conversion[3]]])
    return true_data

def read_data(filename):
    data=[]
    userfile = open(filename, 'r') 
    for n in range(2): line=userfile.readline()
    while line!='':
        temp=[]
        for item in line.split(','): temp.append(item)
        temp[len(temp)-1]=temp[len(temp)-1].split('\n')[0].split('\r')[0]
        data.append(temp)
        line=userfile.readline()
    userfile.close()
    return data

def read_invoice(filename):
    debt=filename.split('/')[-1].split('.docx')[0].split(' - ')[::-1] # Get invoice number and resident name    
    document = Document(filename)
    
    for paragraph in document.paragraphs:
        try: #Make sure date is in the correct format
            if paragraph.text.count("Date:")==1: debt.insert(0, datetime.strptime(paragraph.text.split(': ')[1],'%d %B %Y').strftime('%d/%m/%Y'))
        except:
            print('The date is in the wrong format!\n', 'The date is currently formatted as:', paragraph.text.split(': ')[1],
                      '\n for the following file:\n',  filename+'\n')
            print(datetime.strptime(paragraph.text.split(': ')[1],'%d %B %Y').strftime('%d/%m/%Y'))

    for paragraph in document.tables[0].cell(1,1).paragraphs:
        value='' #Split off pound sign
        for n in range(2,len(list(paragraph.text))): value+=list(paragraph.text)[n]
        
    debt.insert(2,float(value))
    return debt

